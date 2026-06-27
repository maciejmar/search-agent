import io
import re
from dataclasses import dataclass
from pathlib import Path


SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.doc', '.txt'}


@dataclass
class ExtractedDocument:
    name: str
    file_type: str
    text: str


def extract_document_text(filename: str, payload: bytes) -> ExtractedDocument:
    suffix = Path(filename).suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        raise ValueError(f'Unsupported file type: {suffix or "unknown"}')

    if suffix == '.pdf':
        text = _extract_pdf_text(payload)
    elif suffix == '.docx':
        text = _extract_docx_text(payload)
    elif suffix == '.doc':
        text = _extract_legacy_doc_text(payload)
    else:
        text = _extract_plain_text(payload)

    cleaned = normalize_text(text)
    return ExtractedDocument(name=filename, file_type=suffix.lstrip('.'), text=cleaned)


def normalize_text(text: str) -> str:
    lines = []
    for raw_line in re.split(r'\r?\n+', text or ''):
        collapsed = re.sub(r'\s+', ' ', raw_line).strip()
        if collapsed:
            lines.append(collapsed)
    return '\n'.join(lines)


def _extract_pdf_text(payload: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(payload))
    parts: list[str] = []
    for page in reader.pages:
        parts.append(page.extract_text() or '')
    return '\n'.join(parts)


def _extract_docx_text(payload: bytes) -> str:
    from docx import Document

    document = Document(io.BytesIO(payload))
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append(' | '.join(cell.text for cell in row.cells))
    return '\n'.join(parts)


def _extract_plain_text(payload: bytes) -> str:
    for encoding in ('utf-8', 'cp1250', 'latin-1'):
        try:
            return payload.decode(encoding)
        except UnicodeDecodeError:
            continue
    return payload.decode('utf-8', errors='ignore')


def _extract_legacy_doc_text(payload: bytes) -> str:
    decoded = payload.decode('latin-1', errors='ignore')
    fragments = re.findall(r'[A-Za-z0-9_./,:;()@\\/-\\n\\r\\t ]{4,}', decoded)
    return ' '.join(fragments)
